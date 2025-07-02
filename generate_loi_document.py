import sqlite3
import json
import os
from sentence_transformers import SentenceTransformer, util
import torch
import numpy as np
from typing import Dict, List, Optional
import google.generativeai as genai
from docx import Document

# --- CONFIGURATION ---
DB_FILE = "grant_matches.db"
EMBEDDINGS_FILE = "grant_embeddings.json"
MODEL_NAME = 'all-MiniLM-L6-v2'
TOP_N_SIMILAR_GRANTS = 5

# --- IMPORTANT: ADD YOUR GEMINI API KEY HERE ---
# You can get a free key from Google AI Studio: https://aistudio.google.com/
GEMINI_API_KEY = "YOUR_API_KEY"
genai.configure(api_key=GEMINI_API_KEY)

# --- HELPER FUNCTIONS ---

def get_foundation_details(cursor, foundation_ein: str) -> Optional[Dict]:
    """Fetches all necessary details for a foundation."""
    cursor.execute("SELECT name, mission_statement, address_line_1, city, state, zip_code FROM foundations WHERE ein = ?", (foundation_ein,))
    return cursor.fetchone()

def find_best_contact(cursor, foundation_ein: str) -> Optional[str]:
    """Finds the most likely contact person for a foundation."""
    title_priority = ['PRESIDENT', 'EXECUTIVE DIRECTOR', 'DIRECTOR', 'CEO', 'SECRETARY', 'TREASURER', 'TRUSTEE']
    cursor.execute("SELECT name, title FROM officers WHERE foundation_ein = ?", (foundation_ein,))
    officers = cursor.fetchall()
    if not officers: return None
    for title in title_priority:
        for officer in officers:
            if title in (officer['title'] or '').upper():
                return f"{officer['name'].title()}, {officer['title'].title()}"
    return f"{officers[0]['name'].title()}, {officers[0]['title'].title()}"

def get_similar_grants_and_ask_amount(cursor, model, target_foundation_ein: str, charity_profile_text: str) -> (str, str):
    """Finds similar grants and calculates a smart ask amount."""
    cursor.execute("SELECT id, enriched_purpose, grant_amount FROM grants WHERE foundation_ein = ?", (target_foundation_ein,))
    foundation_grants = cursor.fetchall()
    if not foundation_grants:
        return "No similar grants found for this foundation.", "$50,000 (Default)"

    with open(EMBEDDINGS_FILE, 'r') as f: data = json.load(f)
    grant_ids = data['grant_ids']
    grant_embeddings = torch.tensor(data['embeddings'])

    foundation_grant_ids = [g['id'] for g in foundation_grants]
    indices_in_main_list = [grant_ids.index(gid) for gid in foundation_grant_ids if gid in grant_ids]
    if not indices_in_main_list:
        return "Could not find embeddings for this foundation's grants.", "$50,000 (Default)"

    foundation_embeddings = grant_embeddings[indices_in_main_list]
    query_embedding = model.encode(charity_profile_text, convert_to_tensor=True)
    cosine_scores = util.cos_sim(query_embedding, foundation_embeddings)
    top_results = torch.topk(cosine_scores, k=min(TOP_N_SIMILAR_GRANTS, len(foundation_grants)))

    similar_grants_context = "Examples of this foundation's past giving that align with our mission include:\n"
    ask_amounts = []
    for _, idx in zip(top_results[0][0], top_results[1][0]):
        grant_info = foundation_grants[idx]
        if grant_info:
            amount = float(grant_info['grant_amount'] or 0)
            if amount > 0: ask_amounts.append(amount)
            similar_grants_context += f"- A grant of ${amount:,.2f} for: {grant_info['enriched_purpose']}\n"
    
    if ask_amounts:
        avg_amount = np.mean(ask_amounts)
        smart_ask = round(avg_amount / 500) * 500 if avg_amount < 10000 else round(avg_amount / 1000) * 1000
        smart_ask_amount = f"${smart_ask:,.0f}"
    else:
        smart_ask_amount = "$50,000 (Default)"
        
    return similar_grants_context, smart_ask_amount

def generate_loi_paragraph(charity_profile: str, foundation_mission: str, similar_grants_context: str) -> str:
    """Uses the Gemini AI to generate a personalized paragraph for the LOI."""
    print("\n--- Generating AI-Powered Paragraph ---")
    prompt = f"""
    You are an expert grant writer creating a Letter of Inquiry (LOI). Your task is to write a single, compelling paragraph that connects a charity's mission to a foundation's past giving.
    **Charity's Profile:** {charity_profile}
    **Foundation's Mission Statement:** {foundation_mission}
    **Examples of the Foundation's Relevant Past Grants:** {similar_grants_context}
    Based on all of this information, write one concise and persuasive paragraph for the body of the LOI. Start by acknowledging the foundation's history of supporting similar causes, then briefly introduce the charity's project, and explain why the charity's work is a perfect alignment with the foundation's giving patterns. Do not include a greeting or a request for a specific amount. Just write the body paragraph.
    """
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)
        print("AI paragraph generated successfully.")
        return response.text
    except Exception as e:
        print(f"Error generating AI content: {e}")
        return "[An error occurred during AI generation.]"

def create_word_document(report: Dict):
    """Creates a .docx file from the gathered intelligence."""
    doc = Document()
    doc.add_paragraph(f"{report['foundation_name']}\n{report['full_address']}\n")
    doc.add_paragraph(f"Dear {report['contact_person']},")
    doc.add_paragraph("\n")
    doc.add_paragraph(report['ai_paragraph'])
    doc.add_paragraph("\n")
    doc.add_paragraph(f"We are writing to respectfully request a grant of {report['ask_amount']} to support this vital work.")
    doc.add_paragraph("\nSincerely,\n\n[Your Name/Title]\n[Your Organization's Name]")
    
    filename = f"LOI_for_{report['foundation_name'].replace(' ', '_')}.docx"
    doc.save(filename)
    print(f"\n--- Success! Saved Letter of Inquiry to '{filename}' ---")

def main():
    """
    Main function to gather intelligence, generate AI content, and create a Word document.
    """
    if not all(os.path.exists(f) for f in [DB_FILE, EMBEDDINGS_FILE]):
        print(f"Error: Make sure '{DB_FILE}' and '{EMBEDDINGS_FILE}' exist.")
        return

    if GEMINI_API_KEY == "YOUR_API_KEY":
        print("CRITICAL ERROR: Please replace 'YOUR_API_KEY' in the script (line 21) with your actual Gemini API key.")
        return

    print("--- AI Grant Writer - Document Generator ---")

    charity_profile_text = (
        "Our organization, the Urban Youth Arts Collective, serves at-risk youth ages 12-18 in the downtown metropolitan area. "
        "We are seeking $75,000 in program-specific funding for our 'Music & Mentorship' program. "
        "This program provides free, daily after-school music lessons, group workshops, and performance opportunities to 50 students. "
        "Our goal is to improve high school graduation rates by 15% among participants by providing a safe, creative, and supportive environment."
    )
    print("\nUsing Charity Profile: Urban Youth Arts Collective")

    print("\nEnter the EIN of the foundation you want to apply to:")
    target_foundation_ein = input("> ")

    if not target_foundation_ein:
        print("No EIN provided. Exiting.")
        return

    # --- Gather Intelligence ---
    model = SentenceTransformer(MODEL_NAME)
    
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()

    foundation_details = get_foundation_details(cursor, target_foundation_ein)
    if not foundation_details:
        print(f"Error: No foundation found with EIN {target_foundation_ein}.")
        conn.close()
        return

    contact_person = find_best_contact(cursor, target_foundation_ein)
    context, ask_amount = get_similar_grants_and_ask_amount(cursor, model, target_foundation_ein, charity_profile_text)
    
    # --- Generate AI Content ---
    ai_paragraph = generate_loi_paragraph(
        charity_profile_text,
        foundation_details['mission_statement'],
        context
    )

    # --- Assemble and Save Document ---
    report = {
        "foundation_name": foundation_details['name'],
        "full_address": f"{foundation_details['address_line_1']}, {foundation_details['city']}, {foundation_details['state']} {foundation_details['zip_code']}",
        "contact_person": contact_person or "Board of Directors",
        "ask_amount": ask_amount,
        "ai_paragraph": ai_paragraph
    }
    
    create_word_document(report)
    
    conn.close()

if __name__ == "__main__":
    main()
